#!/usr/bin/env python3
from __future__ import annotations

"""
Thesis docx fixes (fact-based + readable).

- Insert "目次" + Word TOC field before Chapter 1.
- Remove the DINO subsection (not executed) and renumber the next heading.
- Replace/insert figures with crisp PNGs generated from local data/logs.

Note: Word needs "Update field" to render the TOC after opening the docx.
"""

from pathlib import Path
import re

import h5py
import numpy as np
import pandas as pd
from PIL import Image, ImageDraw, ImageFont

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm
from docx.text.paragraph import Paragraph


ROOT = Path(".")
DOCX_PATH = ROOT / "学位論文(雛形）.docx"
DATASET_PATH = ROOT / "data" / "dataset.mat"
LABELS_CSV = ROOT / "annotations" / "dataset_labels.csv"
RUN2_REPORT = ROOT / "outputs" / "20260129" / "run_2" / "fastap_report.md"
RUN2_TOPK_INT = ROOT / "outputs" / "20260129" / "run_2" / "topk_internal_query0.png"
RUN2_TOPK_EXT = ROOT / "outputs" / "20260129" / "run_2" / "topk_external_query.png"
FIG_DIR = ROOT / "outputs" / "thesis_figures"


W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"


def ensure_dir(p: Path) -> None:
    p.mkdir(parents=True, exist_ok=True)


def bmode_normalize(frame_2d: np.ndarray) -> np.ndarray:
    a = np.abs(frame_2d.astype(np.float64))
    amax = float(a.max()) if a.size else 1.0
    amax = max(amax, 1e-12)
    x = a / amax
    I = 20.0 * np.log10(np.maximum(x, 1e-12))
    I = np.clip(I, -60.0, 0.0)
    return (I + 60.0) / 60.0


def load_aspect_ratio_from_mat(f: h5py.File) -> float:
    xrefs = f["Acq/x"][...]
    z = np.array(f[xrefs[0, 0]]).reshape(-1)
    x = np.array(f[xrefs[1, 0]]).reshape(-1)
    dz = (float(z.max()) - float(z.min())) / max(1, (len(z) - 1))
    dx = (float(x.max()) - float(x.min())) / max(1, (len(x) - 1))
    return (dz / dx) if dx > 0 else 1.0


def load_frame(idx: int) -> tuple[np.ndarray, float]:
    with h5py.File(DATASET_PATH, "r") as f:
        amp = f["Acq/Amp"]
        fr = np.array(amp[idx], dtype=np.float64)  # (256, 3319)
        r = load_aspect_ratio_from_mat(f)
    return fr.T, r  # (3319, 256)


def to_pil_gray01(img01: np.ndarray) -> Image.Image:
    arr = np.clip(img01 * 255.0, 0, 255).astype(np.uint8)
    return Image.fromarray(arr, mode="L")


def aspect_correct(pil_img: Image.Image, r: float) -> Image.Image:
    w, h = pil_img.size
    new_h = max(1, int(round(h * float(r))))
    return pil_img.resize((w, new_h), Image.Resampling.BICUBIC)


def insert_paragraph_after(paragraph, text: str = "", style: str | None = None):
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if style is not None:
        new_para.style = style
    if text:
        new_para.add_run(text)
    return new_para


def clear_paragraph(paragraph) -> None:
    p = paragraph._p
    for child in list(p):
        p.remove(child)


def paragraph_has_drawing(paragraph) -> bool:
    return bool(paragraph._p.findall(f".//{{{W_NS}}}drawing"))


def find_heading_index(doc: Document, prefix: str) -> int | None:
    for i, p in enumerate(doc.paragraphs):
        if p.style.name.startswith("Heading") and p.text.strip().startswith(prefix):
            return i
    return None


def insert_toc_before_ch1(doc: Document) -> None:
    ch1_idx = find_heading_index(doc, "第 1 章")
    if ch1_idx is None:
        return
    # already inserted?
    for p in doc.paragraphs[:ch1_idx]:
        if p.text.strip() == "目次":
            return
    anchor = doc.paragraphs[ch1_idx]
    title = anchor.insert_paragraph_before("目次", style="Heading 1")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    toc_p = anchor.insert_paragraph_before("", style="Normal")
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), 'TOC \\\\o \"1-3\" \\\\h \\\\z \\\\u')
    toc_p._p.append(fld)

    pb = anchor.insert_paragraph_before("", style="Normal")
    pb.add_run().add_break(WD_BREAK.PAGE)


def remove_dino_section(doc: Document) -> None:
    # Remove from the DINO heading to just before the next 3.1.4 heading, then renumber 3.1.4 -> 3.1.3.
    start = None
    end = None
    for i, p in enumerate(doc.paragraphs):
        if p.style.name.startswith("Heading") and "DINO" in p.text:
            start = i
            continue
        if start is not None and p.style.name.startswith("Heading") and p.text.strip().startswith("3.1.4"):
            end = i
            break
    if start is None:
        return
    if end is None:
        end = start + 1
    for _ in range(end - start):
        p = doc.paragraphs[start]
        p._element.getparent().remove(p._element)
    for p in doc.paragraphs:
        if p.style.name.startswith("Heading") and p.text.strip().startswith("3.1.4"):
            p.text = p.text.replace("3.1.4", "3.1.3", 1)
            break


def draw_line_plot(
    xs: list[float],
    ys: list[float],
    out_path: Path,
    title: str,
    x_label: str,
    y_label: str,
    *,
    width: int = 1600,
    height: int = 900,
) -> None:
    img = Image.new("RGB", (width, height), (255, 255, 255))
    draw = ImageDraw.Draw(img)
    font = ImageFont.load_default()

    left, right, top, bottom = 90, 40, 60, 90
    pw = width - left - right
    ph = height - top - bottom
    draw.rectangle([left, top, left + pw, top + ph], outline=(0, 0, 0), width=2)

    draw.text((left, 20), title[:90], fill=(0, 0, 0), font=font)
    draw.text((left, height - 30), x_label[:60], fill=(0, 0, 0), font=font)
    draw.text((10, top), y_label[:60], fill=(0, 0, 0), font=font)

    if not xs or not ys or len(xs) != len(ys):
        img.save(out_path)
        return

    xmin, xmax = float(min(xs)), float(max(xs))
    ymin, ymax = float(min(ys)), float(max(ys))
    if xmax == xmin:
        xmax = xmin + 1.0
    if ymax == ymin:
        ymax = ymin + 1.0

    def xpx(x: float) -> int:
        return int(round(left + (x - xmin) / (xmax - xmin) * pw))

    def ypx(y: float) -> int:
        return int(round(top + (1.0 - (y - ymin) / (ymax - ymin)) * ph))

    pts = [(xpx(float(x)), ypx(float(y))) for x, y in zip(xs, ys)]
    if len(pts) >= 2:
        draw.line(pts, fill=(30, 80, 200), width=3)
    for px, py in pts:
        draw.ellipse([px - 3, py - 3, px + 3, py + 3], fill=(30, 80, 200))
    img.save(out_path)


def draw_pseudo_label_figure(out_path: Path) -> dict[str, float]:
    df = pd.read_csv(LABELS_CSV).sort_values("frame")
    frames = df["frame"].astype(int).to_list()
    labels = df["label"].astype(int).to_list()
    smooth = df["smooth_score"].astype(float).to_list()

    thr = float(np.percentile(np.asarray(smooth, dtype=np.float64), 20.0))
    pos_n = int(np.sum(np.asarray(labels, dtype=np.int32) == 1))
    pos_ratio = float(pos_n / max(1, len(labels)))

    width, height = 1800, 950
    img = Image.new("RGB", (width, height), (255, 255, 255))
    draw = ImageDraw.Draw(img)
    font = ImageFont.load_default()

    draw.text((20, 15), "Pseudo labels from dataset_labels.csv", fill=(0, 0, 0), font=font)
    draw.text((20, 35), f"N={len(frames)}  pos={pos_n} ({pos_ratio*100:.1f}%)  thr(p20)={thr:.3f}", fill=(0, 0, 0), font=font)

    left, right, top, mid = 80, 40, 70, 650
    pw = width - left - right
    ph = mid - top
    draw.rectangle([left, top, left + pw, top + ph], outline=(0, 0, 0), width=2)

    xmin, xmax = min(frames), max(frames)
    ymin, ymax = min(smooth), max(smooth)
    if xmax == xmin:
        xmax = xmin + 1
    if ymax == ymin:
        ymax = ymin + 1.0

    def xpx(x: int) -> int:
        return int(round(left + (x - xmin) / (xmax - xmin) * pw))

    def ypx(y: float) -> int:
        return int(round(top + (1.0 - (y - ymin) / (ymax - ymin)) * ph))

    pts = [(xpx(x), ypx(y)) for x, y in zip(frames, smooth)]
    if len(pts) >= 2:
        draw.line(pts, fill=(50, 50, 50), width=2)
    y_thr = ypx(thr)
    draw.line([(left, y_thr), (left + pw, y_thr)], fill=(200, 50, 50), width=2)
    draw.text((left + 6, max(top, y_thr - 16)), "p20 threshold", fill=(200, 50, 50), font=font)

    bar_top = mid + 40
    bar_h = 70
    draw.rectangle([left, bar_top, left + pw, bar_top + bar_h], outline=(0, 0, 0), width=2)
    for x, y in zip(frames, labels):
        px = xpx(x)
        color = (60, 160, 80) if y == 1 else (220, 220, 220)
        draw.line([(px, bar_top + 1), (px, bar_top + bar_h - 1)], fill=color, width=2)
    draw.text((left, bar_top + bar_h + 10), "label=1 (green), label=0 (gray)", fill=(0, 0, 0), font=font)

    img.save(out_path)
    return {"pos_ratio": pos_ratio, "pos_n": float(pos_n)}


def parse_simsiam_losses(md_text: str) -> tuple[list[int], list[float]]:
    epochs: list[int] = []
    losses: list[float] = []
    in_table = False
    for line in md_text.splitlines():
        if line.strip() == "## SimSiam Training History":
            in_table = False
            continue
        if line.startswith("| Epoch | Loss |"):
            in_table = True
            continue
        if in_table:
            if not line.startswith("|"):
                break
            cols = [c.strip() for c in line.strip().strip("|").split("|")]
            if len(cols) < 2 or cols[0] in ("Epoch", "-----"):
                continue
            try:
                epochs.append(int(cols[0]))
                losses.append(float(cols[1]))
            except Exception:
                continue
    return epochs, losses


def replace_figure_paragraph(doc: Document, para_idx: int, img_path: Path, width_cm: float) -> None:
    p = doc.paragraphs[para_idx]
    clear_paragraph(p)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(img_path), width=Cm(width_cm))


def main() -> int:
    ensure_dir(FIG_DIR)
    if not (DOCX_PATH.exists() and DATASET_PATH.exists() and LABELS_CSV.exists() and RUN2_REPORT.exists()):
        raise SystemExit("missing required inputs (docx/dataset/labels/run_2 report)")

    # Generate crisp figures
    raw0, r = load_frame(0)
    fig_2_1 = FIG_DIR / "fig_2_1_bmode_example.png"
    img0 = aspect_correct(to_pil_gray01(bmode_normalize(raw0)), r)
    img0 = img0.resize((1400, int(round(1400 * img0.size[1] / img0.size[0]))), Image.Resampling.BICUBIC)
    img0.save(fig_2_1)

    fig_2_2 = FIG_DIR / "fig_2_2_pseudo_labels.png"
    label_stats = draw_pseudo_label_figure(fig_2_2)

    epochs, losses = parse_simsiam_losses(RUN2_REPORT.read_text(encoding="utf-8", errors="ignore"))
    fig_3_1 = FIG_DIR / "fig_3_1_simsiam_loss.png"
    draw_line_plot([float(e) for e in epochs], [float(l) for l in losses], fig_3_1, "SimSiam loss (run_2 log)", "epoch", "loss")

    fig_4_int = FIG_DIR / "fig_4_1_topk_internal.png"
    fig_4_ext = FIG_DIR / "fig_4_2_topk_external.png"
    Image.open(RUN2_TOPK_INT).save(fig_4_int)
    Image.open(RUN2_TOPK_EXT).save(fig_4_ext)

    # Update docx
    doc = Document(str(DOCX_PATH))
    insert_toc_before_ch1(doc)
    remove_dino_section(doc)

    # 2.2.1: replace placeholder drawing paragraph
    idx_221 = find_heading_index(doc, "2.2.1")
    if idx_221 is not None:
        draw_idx = None
        for j in range(idx_221 + 1, min(idx_221 + 6, len(doc.paragraphs))):
            if paragraph_has_drawing(doc.paragraphs[j]):
                draw_idx = j
                break
        if draw_idx is not None:
            replace_figure_paragraph(doc, draw_idx, fig_2_1, width_cm=14.5)
            # Caption line right after
            for j in range(draw_idx + 1, min(draw_idx + 4, len(doc.paragraphs))):
                if doc.paragraphs[j].style.name == "Normal" and doc.paragraphs[j].text.strip():
                    doc.paragraphs[j].text = "図 2.1  B-mode 正規化および縦横比補正後のフレーム例"
                    doc.paragraphs[j].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    break

    # 2.2.2: make text factual + insert figure
    idx_222 = find_heading_index(doc, "2.2.2")
    if idx_222 is not None:
        desc_idx = None
        for j in range(idx_222 + 1, min(idx_222 + 6, len(doc.paragraphs))):
            if doc.paragraphs[j].style.name == "Normal" and doc.paragraphs[j].text.strip():
                desc_idx = j
                break
        if desc_idx is not None:
            pos_n = int(label_stats["pos_n"])
            pos_ratio = float(label_stats["pos_ratio"])
            doc.paragraphs[desc_idx].text = f"dataset_labels.csv により安定フレームを 0/1 で付与し、陽性は {pos_n} 枚（{pos_ratio*100:.1f}%）であった。"
            pic_p = insert_paragraph_after(doc.paragraphs[desc_idx], "", style="Normal")
            pic_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            pic_p.add_run().add_picture(str(fig_2_2), width=Cm(15.5))
            cap_p = insert_paragraph_after(pic_p, "図 2.2  擬似ラベル生成に用いた smooth_score と閾値、および 0/1 ラベル", style="Normal")
            cap_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 3.2.2: replace placeholder drawing paragraph
    idx_322 = find_heading_index(doc, "3.2.2")
    if idx_322 is not None:
        draw_idx = None
        for j in range(idx_322 + 1, min(idx_322 + 6, len(doc.paragraphs))):
            if paragraph_has_drawing(doc.paragraphs[j]):
                draw_idx = j
                break
        if draw_idx is not None:
            replace_figure_paragraph(doc, draw_idx, fig_3_1, width_cm=14.5)
            for j in range(draw_idx + 1, min(draw_idx + 4, len(doc.paragraphs))):
                if doc.paragraphs[j].style.name == "Normal" and doc.paragraphs[j].text.strip():
                    doc.paragraphs[j].text = "図 3.1  SimSiam 事前学習の損失推移（学習ログより）"
                    doc.paragraphs[j].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    break

    # 4.2.2: insert internal top-k
    idx_422 = find_heading_index(doc, "4.2.2")
    if idx_422 is not None:
        desc_idx = None
        for j in range(idx_422 + 1, min(idx_422 + 8, len(doc.paragraphs))):
            if doc.paragraphs[j].style.name == "Normal" and doc.paragraphs[j].text.strip():
                desc_idx = j
                break
        if desc_idx is not None:
            pic_p = insert_paragraph_after(doc.paragraphs[desc_idx], "", style="Normal")
            pic_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            pic_p.add_run().add_picture(str(fig_4_int), width=Cm(16.0))
            cap_p = insert_paragraph_after(pic_p, "図 4.1  内部クエリ（index=0）の Top-10 可視化", style="Normal")
            cap_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 4.2.3: replace placeholder drawing paragraph with external top-k
    idx_423 = find_heading_index(doc, "4.2.3")
    if idx_423 is not None:
        draw_idx = None
        for j in range(idx_423 + 1, min(idx_423 + 10, len(doc.paragraphs))):
            if paragraph_has_drawing(doc.paragraphs[j]):
                draw_idx = j
                break
        if draw_idx is not None:
            replace_figure_paragraph(doc, draw_idx, fig_4_ext, width_cm=16.0)
            for j in range(draw_idx + 1, min(draw_idx + 4, len(doc.paragraphs))):
                if doc.paragraphs[j].style.name == "Normal" and doc.paragraphs[j].text.strip():
                    doc.paragraphs[j].text = "図 4.2  外部参照画像の Top-10 可視化"
                    doc.paragraphs[j].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    break

    doc.save(str(DOCX_PATH))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())

